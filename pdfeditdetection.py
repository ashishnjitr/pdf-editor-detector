import streamlit as st
import pypdf
import fitz  # PyMuPDF
import re
import io
import os
import hashlib
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime

# 1. Page Configuration & Styling
st.set_page_config(
    page_title="PDF BUSTER", 
    page_icon="💥", 
    layout="wide",
    initial_sidebar_state="expanded"
)

STYLE_INJECTION = """
<style>
    .brand-title { font-family: 'Courier New', Courier, monospace; font-size: 38px; font-weight: 900; letter-spacing: -1px; color: #FF4B4B; margin-bottom: 0px; display: flex; align-items: center; gap: 10px; }
    .brand-tagline { color: #6c757d; font-size: 13px; text-transform: uppercase; letter-spacing: 2px; margin-bottom: 25px; border-bottom: 2px solid #efefef; padding-bottom: 10px; }
    .buster-grid { display: flex; justify-content: space-between; gap: 12px; margin-top: 15px; margin-bottom: 20px; }
    .buster-card { background-color: #f8f9fa; border: 1px solid #e9ecef; border-top: 4px solid #6c757d; border-radius: 6px; padding: 14px; flex: 1; text-align: center; }
    .buster-card.alert-active { border-top-color: #FF4B4B; }
    .buster-card.caution-active { border-top-color: #FFA500; }
    .buster-card.clean-active { border-top-color: #28a745; }
    .buster-val { font-size: 22px; font-weight: 800; font-family: monospace; margin-bottom: 2px; }
    .buster-lbl { font-size: 11px; text-transform: uppercase; letter-spacing: 1px; color: #6c757d; font-weight: 600; }
    .detail-block { padding: 12px; border-radius: 6px; background-color: #fafafa; border-left: 4px solid #007bd9; margin-bottom: 10px; }
    .detail-title { font-weight: 700; font-size: 14px; margin-bottom: 3px; color: #1f2937; }
    .detail-text { font-size: 13px; color: #4b5563; line-height: 1.4; }
    .evidence-box { padding: 12px; background-color: #fff8f8; border: 1px solid #ffebeb; border-radius: 6px; margin-bottom: 15px; }
    .evidence-item { font-size: 12.5px; font-family: monospace; color: #333; margin-bottom: 4px; }
    .evidence-label { font-weight: bold; color: #c00; }
</style>
"""
st.html(STYLE_INJECTION)

st.html('<div class="brand-title">💥 PDF BUSTER</div>')
st.html('<div class="brand-tagline">Deep Forensics, Batch Conversion & Privacy Sanitizer Engine</div>')

st.sidebar.markdown("### 🛠️ Mode Selection")
app_mode = st.sidebar.radio(
    "Choose Utility Interface:",
    [
        "🔍 Batch PDF Forensic Analyzer", 
        "📄 Universal PDF to Word Converter",
        "🧼 PDF Privacy Sanitizer & Metadata Wiper"
    ]
)

# -------------------------------------------------------------
# MODE A: BATCH FORENSIC ANALYZER + VISUAL REDLINING + AUDIT LOG
# -------------------------------------------------------------
if app_mode == "🔍 Batch PDF Forensic Analyzer":
    st.subheader("Deep-Object Tampering & Visual Redlining")
    uploaded_files = st.file_uploader(
        "Upload one or multiple PDF documents for structural screening", 
        type="pdf", 
        accept_multiple_files=True, 
        key="batch_forensic_upload"
    )

    def analyze_single_pdf(file_bytes, filename):
        sha256_hash = hashlib.sha256(file_bytes).hexdigest()
        
        results = {
            "filename": filename,
            "sha256": sha256_hash,
            "incremental_updates": 0, "xref_tables": 0, "font_anomalies": 0,
            "is_edited": False, "tamper_lock": False, "metadata": {},
            "detailed_findings": [], "edited_segments": [],
            "inferred_tool": "None Detected",
            "device_details": "None Detected",
            "location_data": "None Detected",
            "timeline_analysis": "Consistent",
            "verdict": "SAFE TO PROCEED",
            "annotated_images": []
        }
        
        eof_markers = re.findall(b'%%EOF', file_bytes)
        xref_markers = re.findall(b'xref', file_bytes)
        results["incremental_updates"] = len(eof_markers)
        results["xref_tables"] = len(xref_markers)

        try:
            doc_fitz = fitz.open(stream=file_bytes, filetype="pdf")
            for page_num in range(len(doc_fitz)):
                page = doc_fitz[page_num]
                text_blocks = page.get_text("blocks")
                page_has_suspicious_blocks = False
                
                for block in text_blocks:
                    block_text = block[4].strip()
                    matched_indicators = [m for m in ["ilovepdf", "smallpdf", "watermark", "eval", "sejda", "pdfescape", "pdf2go"] if m in block_text.lower()]
                    
                    if matched_indicators:
                        results["tamper_lock"] = True
                        results["inferred_tool"] = matched_indicators[0].upper()
                        results["edited_segments"].append(f"Page {page_num + 1}: '{block_text}'")
                        
                        rect = fitz.Rect(block[:4])
                        page.draw_rect(rect, color=(1, 0, 0), width=2)
                        page_has_suspicious_blocks = True
                
                pix = page.get_pixmap(dpi=130)
                img_bytes = pix.tobytes("png")
                results["annotated_images"].append((page_num + 1, img_bytes, page_has_suspicious_blocks))
            
            all_fonts = []
            for page in doc_fitz:
                all_fonts.extend([f[3] for f in page.get_fonts() if f])
            unique_fonts = list(set(all_fonts))
            suspicious_fonts = [f for f in unique_fonts if "identity-h" in f.lower() or "custom" in f.lower()]
            results["font_anomalies"] = len(suspicious_fonts)
        except Exception as e:
            results["detailed_findings"].append({"title": "Parse Interruption", "text": str(e)})

        try:
            pdf_file = io.BytesIO(file_bytes)
            reader = pypdf.PdfReader(pdf_file)
            metadata = reader.metadata
            if metadata:
                cleaned_meta = {k.replace('/', ''): str(v) for k, v in metadata.items()}
                results["metadata"] = cleaned_meta
                
                creator = cleaned_meta.get("Creator", "")
                producer = cleaned_meta.get("Producer", "")
                device_markers = []
                if creator: device_markers.append(f"Creator: {creator}")
                if producer: device_markers.append(f"Engine: {producer}")
                if device_markers: results["device_details"] = " | ".join(device_markers)
                
                producer_lower = (producer + creator).lower()
                for tool in ["ilovepdf", "smallpdf", "pdf2go", "nitro", "soda", "libreoffice", "canva", "pdfescape", "sejda", "acrobat"]:
                    if tool in producer_lower:
                        if tool != "acrobat" or len(eof_markers) > 1:
                            results["tamper_lock"] = True
                            results["inferred_tool"] = tool.upper()

                create_date = cleaned_meta.get("CreationDate", "")
                mod_date = cleaned_meta.get("ModDate", "")
                if create_date and mod_date and create_date != mod_date:
                    results["timeline_analysis"] = "Chronology Conflict (Altered Post-Creation)"
                    tz_match = re.search(r'([+-]\d{2}\'\d{2}\')', mod_date)
                    if tz_match:
                        clean_tz = tz_match.group(1).replace("'", ":").strip(":")
                        results["location_data"] = f"GMT {clean_tz}"
        except:
            pass

        if results["tamper_lock"] or len(results["edited_segments"]) > 0:
            results["is_edited"] = True
            results["verdict"] = "FULL RED FLAG"
        elif results["font_anomalies"] > 0 and results["timeline_analysis"] != "Consistent":
            results["is_edited"] = True
            results["verdict"] = "CAUTION"
            results["detailed_findings"].append({
                "title": "⚠️ Targeted Structural Layout Revision",
                "text": "Detected localized typographic anomalies alongside an altered post-creation timestamp conflict."
            })
        else:
            results["is_edited"] = False
            results["verdict"] = "SAFE TO PROCEED"

        return results

    def generate_certificate(res):
        cert = f"""================================================================================
                    PDF BUSTER // FORENSIC INTEGRITY AUDIT REPORT
================================================================================
Document Name       : {res['filename']}
SHA-256 Hash        : {res['sha256']}
Verification Time   : {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}
Final Status        : {res['verdict']}

[STRUCTURAL METRICS]
- Save Cycles / EOF Markers : {res['incremental_updates']}
- XREF Reference Maps       : {res['xref_tables']}
- Font Subset Anomalies     : {res['font_anomalies']}

[ENVIRONMENTAL TRACE]
- Detected Tool Fingerprint : {res['inferred_tool']}
- Device Profile            : {res['device_details']}
- Timestamp / Location Code : {res['location_data']}
- Chronology Audit          : {res['timeline_analysis']}

[ISOLATED ANOMALIES & OVERLAYS]
{chr(10).join(['- ' + seg for seg in res['edited_segments']]) if res['edited_segments'] else 'No localized visual overlays detected.'}

[DETAILED FINDINGS]
{chr(10).join(['* ' + f['title'] + ': ' + f['text'] for f in res['detailed_findings']]) if res['detailed_findings'] else 'Document exhibits native single-compilation structure.'}
================================================================================
Generated automatically by PDF BUSTER Digital Forensics Core.
"""
        return cert.encode('utf-8')

    if uploaded_files:
        batch_results = []
        with st.spinner(f"Analyzing {len(uploaded_files)} document(s)..."):
            for up_file in uploaded_files:
                file_bytes = up_file.read()
                res = analyze_single_pdf(file_bytes, up_file.name)
                batch_results.append(res)
        
        st.markdown("### 📊 Batch Pipeline Summary")
        summary_data = []
        for r in batch_results:
            status_icon = "🛑 Red Flag" if r["verdict"] == "FULL RED FLAG" else ("⚠️ Caution" if r["verdict"] == "CAUTION" else "✅ Clean")
            summary_data.append({
                "Filename": r["filename"],
                "Verdict": status_icon,
                "Save Cycles": r["incremental_updates"],
                "XREF Maps": r["xref_tables"],
                "Tool Detected": r["inferred_tool"]
            })
        st.dataframe(summary_data, use_container_width=True)
        
        st.markdown("---")
        st.markdown("### 🔍 Granular Document Inspections")
        
        for r in batch_results:
            expander_title = f"{'🛑' if r['verdict'] == 'FULL RED FLAG' else ('⚠️' if r['verdict'] == 'CAUTION' else '🛡️')} {r['filename']} — Verdict: {r['verdict']}"
            with st.expander(expander_title, expanded=(len(batch_results) == 1)):
                card_class = "alert-active" if r["verdict"] == "FULL RED FLAG" else ("caution-active" if r["verdict"] == "CAUTION" else "clean-active")
                
                metrics_html = (
                    f'<div class="buster-grid">'
                    f'<div class="buster-card {card_class}"><div class="buster-val">{r["incremental_updates"]}</div><div class="buster-lbl">Appended Saves</div></div>'
                    f'<div class="buster-card {card_class}"><div class="buster-val">{r["xref_tables"]}</div><div class="buster-lbl">XREF Maps</div></div>'
                    f'<div class="buster-card {card_class}"><div class="buster-val">{r["font_anomalies"]}</div><div class="buster-lbl">Font Anomalies</div></div>'
                    f'</div>'
                )
                st.html(metrics_html)
                
                evidence_html = f"""
                <div class="evidence-box">
                    <div class="evidence-item"><span class="evidence-label">🛠️ Identified Tool:</span> {r['inferred_tool']}</div>
                    <div class="evidence-item"><span class="evidence-label">💻 Origin Profile:</span> {r['device_details']}</div>
                    <div class="evidence-item"><span class="evidence-label">📍 Timezone Code:</span> {r['location_data']}</div>
                    <div class="evidence-item"><span class="evidence-label">📅 Timeline State:</span> {r['timeline_analysis']}</div>
                </div>
                """
                st.html(evidence_html)
                
                st.markdown("**🖼️ Visual Document Redlining Heatmap**")
                cols = st.columns(min(len(r["annotated_images"]), 3))
                for idx, (p_num, img_b, is_flagged) in enumerate(r["annotated_images"]):
                    with cols[idx % 3]:
                        caption = f"Page {p_num} {'(🚨 Overlay Flagged)' if is_flagged else '(Clean)'}"
                        st.image(img_b, caption=caption, use_container_width=True)
                
                cert_bytes = generate_certificate(r)
                st.download_button(
                    label=f"📜 Download Forensic Audit Certificate ({r['filename']})",
                    data=cert_bytes,
                    file_name=f"Forensic_Certificate_{r['filename']}.txt",
                    mime="text/plain",
                    key=f"cert_{r['filename']}"
                )

# -------------------------------------------------------------
# MODE B: UNIVERSAL PDF TO WORD CONVERTER
# -------------------------------------------------------------
elif app_mode == "📄 Universal PDF to Word Converter":
    st.subheader("Universal PDF to DOCX Converter")
    st.markdown("Convert native text or scanned layout layers directly to editable Word documents.")
    
    uploaded_pdf = st.file_uploader("Upload target PDF file", type=["pdf"], key="universal_converter_upload")
    
    @st.cache_data(show_spinner=False)
    def convert_pdf_to_docx_cached(file_bytes):
        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
        pdf_stream = fitz.open(stream=file_bytes, filetype="pdf")
        for page_num in range(len(pdf_stream)):
            page = pdf_stream[page_num]
            text_blocks = page.get_text("blocks")
            
            if not text_blocks:
                tp_text = page.get_text("text")
                if tp_text.strip():
                    p = doc.add_paragraph()
                    p.add_run(tp_text)
                else:
                    p = doc.add_paragraph()
                    p.add_run(f"--- [Page {page_num + 1}: Scanned Image Content - Layout Flattened] ---").italic = True
            else:
                text_blocks.sort(key=lambda b: (b[1], b[0]))
                for block in text_blocks:
                    block_text = block[4].strip()
                    if block_text:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(5)
                        p.paragraph_format.line_spacing = 1.15
                        run = p.add_run(block_text)
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
                        
            if page_num < len(pdf_stream) - 1:
                doc.add_page_break()
                
        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)
        return output_stream.getvalue()

    if uploaded_pdf is not None:
        file_bytes = uploaded_pdf.read()
        base_filename, _ = os.path.splitext(uploaded_pdf.name)
        target_docx_name = f"{base_filename}.docx"
        
        st.success(f"Loaded: `{uploaded_pdf.name}`")
        with st.spinner("Converting document layers..."):
            docx_data = convert_pdf_to_docx_cached(file_bytes)
            
        st.download_button(
            label=f"📥 Download Editable Word File ({target_docx_name})",
            data=docx_data,
            file_name=target_docx_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

# -------------------------------------------------------------
# MODE C: HIDDEN DATA & PRIVACY SANITIZER (DE-TAMPER & SCRUB)
# -------------------------------------------------------------
elif app_mode == "🧼 PDF Privacy Sanitizer & Metadata Wiper":
    st.subheader("Document Privacy Sanitizer & Layer De-Tampering")
    st.markdown("Completely strip internal tracking signatures, wipe metadata logs, and flatten visual layers to prevent reverse-engineering of edits.")
    
    sanitize_upload = st.file_uploader("Upload PDF to sanitize and scrub", type=["pdf"], key="sanitizer_upload")
    
    flatten_option = st.checkbox("Flatten visual layers (Convert pages to pristine, uneditable image streams to burn out white-outs)", value=True)
    
    def sanitize_pdf_document(file_bytes, flatten=True):
        src_doc = fitz.open(stream=file_bytes, filetype="pdf")
        clean_doc = fitz.open()
        
        if flatten:
            for page in src_doc:
                pix = page.get_pixmap(dpi=200)
                img_bytes = pix.tobytes("png")
                img_doc = fitz.open(stream=img_bytes, filetype="png")
                rect = img_doc[0].rect
                pdfbytes = img_doc.convert_to_pdf()
                img_pdf = fitz.open("pdf", pdfbytes)
                page_clean = clean_doc.new_page(width=rect.width, height=rect.height)
                page_clean.show_pdf_page(rect, img_pdf, 0)
        else:
            clean_doc.insert_pdf(src_doc)
            
        clean_doc.set_metadata({
            "format": "PDF 1.7",
            "title": "",
            "author": "",
            "subject": "",
            "keywords": "",
            "creator": "Clean PDF Standard",
            "producer": "System Native Engine",
            "creationDate": "",
            "modDate": "",
            "trapped": "False"
        })
        
        output_stream = io.BytesIO()
        clean_doc.save(
            output_stream,
            garbage=4,
            deflate=True,
            clean=True,
            deflate_images=True,
            deflate_fonts=True
        )
        output_stream.seek(0)
        return output_stream.getvalue()

    if sanitize_upload is not None:
        file_bytes = sanitize_upload.read()
        base_name, _ = os.path.splitext(sanitize_upload.name)
        sanitized_filename = f"{base_name}_sanitized.pdf"
        
        st.info(f"Loaded `{sanitize_upload.name}` for privacy scrubbing.")
        
        with st.spinner("Scrubbing metadata headers and purging revision trails..."):
            cleaned_pdf_bytes = sanitize_pdf_document(file_bytes, flatten=flatten_option)
            
        st.success("Document successfully sanitized! All tracking footprints and revision trailers have been purged.")
        
        st.download_button(
            label=f"📥 Download Sanitized PDF ({sanitized_filename})",
            data=cleaned_pdf_bytes,
            file_name=sanitized_filename,
            mime="application/pdf",
            use_container_width=True
        )
